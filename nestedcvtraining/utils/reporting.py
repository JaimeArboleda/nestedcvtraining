from docx.shared import Inches, Mm, Pt
import numpy as np
from .plotting import plot_calibration_curve, plot_roc_curve, plot_precision_recall_curve, plot_confusion_matrix, plot_histogram
from .metrics import get_metric
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docxcompose.composer import Composer
from collections import Counter
from datetime import datetime


SIZE_IMAGE = 8
SIZE_FOLD_COL = 0.9
SIZE_MODEL_COL = 1.2
SIZE_PARAM_COL = 2.5
SIZE_COMMENT_COL = 3.5
SIZE_METRIC_COL = 1.5
SIZE_SMALL_FONT = 8
SIZE_NORMAL_FONT = 11


class MetadataFit:

    def __init__(
            self,
            num_init_samples_bf,
            prop_minority_class_bf,
            num_init_samples_af=None,
            prop_minority_class_af=None
    ):
        self._num_init_samples_bf = num_init_samples_bf
        self._prop_minority_class_bf = prop_minority_class_bf
        if num_init_samples_af:
            self._num_init_samples_af = num_init_samples_af
        else:
            self._num_init_samples_af = num_init_samples_bf
        if prop_minority_class_af:
            self._prop_minority_class_af = prop_minority_class_af
        else:
            self._prop_minority_class_af = prop_minority_class_bf

    def get_num_init_samples_bf(self):
        return self._num_init_samples_bf

    def get_prop_minority_class_bf(self):
        return self._prop_minority_class_bf

    def get_num_init_samples_af(self):
        return self._num_init_samples_af

    def get_prop_minority_class_af(self):
        return self._prop_minority_class_af


def averaged_metadata_list(metadata_fit_list):

    return MetadataFit(
            num_init_samples_bf=np.mean(
                [metadata_fit.get_num_init_samples_bf() for metadata_fit in metadata_fit_list]),
            prop_minority_class_bf=np.mean(
                [metadata_fit.get_prop_minority_class_bf() for metadata_fit in metadata_fit_list]),
            num_init_samples_af = np.mean(
                [metadata_fit.get_num_init_samples_af() for metadata_fit in metadata_fit_list]),
            prop_minority_class_af = np.mean(
                [metadata_fit.get_prop_minority_class_af() for metadata_fit in metadata_fit_list])
    )


def ordered_keys_classes(binary_counter):
    return [key for key, _ in reversed(binary_counter.most_common())]


def prop_minority_to_rest_class(binary_counter):
    ordered_counter = binary_counter.most_common()
    num_minority_class = ordered_counter[-1][1]
    num_rest = sum([value for value in binary_counter.values()])
    return num_minority_class / num_rest


def color_row(row):
    # Make row of cells background colored
    for cell in row.cells:
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="FFA07A"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_2)


def add_plots_doc(report_doc, ys, y_probas, folds_index):
    report_doc.add_heading(f'Main plots', level=2)

    # Plot calibration curve
    report_doc.add_heading('Calibration plots', level=3)
    for index, y in enumerate(ys):
        report_doc.add_heading(f'Calibration plot of fold {folds_index[index]}', level=4)
        memfile = plot_calibration_curve(y, y_probas[index])
        report_doc.add_picture(memfile, width=Inches(SIZE_IMAGE))
        memfile.close()

    # Plot precision recall curve
    report_doc.add_heading('Precision-recall curve plots', level=3)
    for index, y in enumerate(ys):
        report_doc.add_heading(f'Precision-recall curve plot of fold {folds_index[index]}', level=4)
        memfile = plot_precision_recall_curve(y, y_probas[index])
        report_doc.add_picture(memfile, width=Inches(SIZE_IMAGE))
        memfile.close()

    # Plot roc curve
    report_doc.add_heading('ROC curve plots', level=3)
    for index, y in enumerate(ys):
        report_doc.add_heading(f'ROC curve plot of fold {folds_index[index]}', level=4)
        memfile = plot_roc_curve(y, y_probas[index])
        report_doc.add_picture(memfile, width=Inches(SIZE_IMAGE))
        memfile.close()

    # Plot confussion matrix
    report_doc.add_heading('Confusion matrix', level=3)
    for index, y in enumerate(ys):
        report_doc.add_heading(f'Confusion matrix of fold {folds_index[index]}', level=4)
        memfile = plot_confusion_matrix(y, y_probas[index])
        report_doc.add_picture(memfile, width=Inches(SIZE_IMAGE))
        memfile.close()

    # Plot histogram
    report_doc.add_heading('Histograms', level=3)
    for index, y in enumerate(ys):
        report_doc.add_heading(f'Histogram of fold {folds_index[index]}', level=4)
        memfile = plot_histogram(y, y_probas[index])
        report_doc.add_picture(memfile, width=Inches(SIZE_IMAGE))
        memfile.close()


def evaluate_model(dict_models, Xs, ys, X_val_var, y_val_var, folds_index,
                   loss_metric, peeking_metrics, report_doc, add_plots):
    y_probas = []
    for index, dict_model in enumerate(dict_models):
        y_probas.append(dict_model['model'].predict_proba(Xs[index])[:, 1])

    # Add a table for comparison of metrics
    if peeking_metrics:
        report_doc.add_heading(f'Winner models of each fold and main metrics', level=2)
        table = report_doc.add_table(rows=len(dict_models) + 1, cols=len(peeking_metrics) + 4)
        table.style = 'TableGrid'
        table.autofit = False
        # Fill first row of table
        first_row = table.rows[0]
        first_row.cells[0].width = Inches(SIZE_FOLD_COL)
        first_row.cells[0].paragraphs[0].add_run('Fold').bold = True
        first_row.cells[1].width = Inches(SIZE_MODEL_COL)
        first_row.cells[1].paragraphs[0].add_run('Model').bold = True
        first_row.cells[2].width = Inches(SIZE_PARAM_COL)
        first_row.cells[2].paragraphs[0].add_run('Params').bold = True
        first_row.cells[3].width = Inches(SIZE_COMMENT_COL)
        first_row.cells[3].paragraphs[0].add_run('Comments').bold = True
        for i, metric in enumerate(peeking_metrics):
            first_row.cells[i + 4].width = Inches(SIZE_METRIC_COL)
            first_row.cells[i + 4].paragraphs[0].add_run(metric).bold = True
        # Fill other rows of table
        metrics = []
        for index, dict_model in enumerate(dict_models):
            row_index = index + 1
            row = table.rows[row_index]
            row.cells[0].width = Inches(SIZE_FOLD_COL)
            row.cells[0].text = str(folds_index[index])
            row.cells[1].width = Inches(SIZE_MODEL_COL)
            row.cells[1].text = dict_model['params']['model']
            row.cells[2].width = Inches(SIZE_PARAM_COL)
            write_paragraphs_dict(row.cells[2],
                                  {key: value
                                   for (key, value) in dict_model['params'].items() if key != 'model'},
                                  SIZE_SMALL_FONT,
                                  is_cell_table=True)
            row.cells[3].width = Inches(SIZE_COMMENT_COL)
            write_paragraphs_dict(row.cells[3], dict_model['comments'], SIZE_SMALL_FONT, is_cell_table=True)
            for i, metric in enumerate(peeking_metrics):
                row.cells[i + 4].width = Inches(SIZE_METRIC_COL)
                value_of_metric = get_metric(metric, 'real')(ys[index], y_probas[index])
                if metric == loss_metric:
                    metrics.append(value_of_metric)
                row.cells[i + 4].paragraphs[0].add_run(
                    str(np.round(value_of_metric, 3)))
        # Add average
        report_doc.add_paragraph(f'For the selected optimization metric {loss_metric} '
                                 f'the average score is {np.round(np.mean(metrics), 3)}'
                                 f', and the standard deviation is {np.round(np.std(metrics), 3)}.')

    if add_plots:
        add_plots_doc(report_doc, ys, y_probas, folds_index)

    report_doc.add_heading('Comparison of several predictions to assess variance', level=2)
    table = report_doc.add_table(rows=len(y_val_var) + 1, cols=len(dict_models) + 3)
    table.style = 'TableGrid'
    table.autofit = False
    # Fill first row of table
    first_row = table.rows[0]
    first_row.cells[0].width = Inches(SIZE_FOLD_COL)
    first_row.cells[0].paragraphs[0].add_run('Instance').bold = True
    first_row.cells[1].width = Inches(SIZE_FOLD_COL)
    first_row.cells[1].paragraphs[0].add_run('Real label').bold = True
    for i in range(len(dict_models)):
        first_row.cells[i+2].width = Inches(SIZE_METRIC_COL)
        first_row.cells[i+2].paragraphs[0].add_run(f'Prediction by model of fold {folds_index[i]}').bold = True
    first_row.cells[len(dict_models) + 2].width = Inches(SIZE_METRIC_COL)
    first_row.cells[len(dict_models) + 2].paragraphs[0].add_run('Standard deviation in predictions of this instance').bold = True
    # Fill other rows
    stds = []
    for index, y in enumerate(y_val_var):
        predictions = []
        row_index = index + 1
        row = table.rows[row_index]
        row.cells[0].width = Inches(SIZE_FOLD_COL)
        row.cells[0].paragraphs[0].add_run(str(row_index))
        row.cells[1].width = Inches(SIZE_FOLD_COL)
        row.cells[1].paragraphs[0].add_run(str(y))
        for i, dict_model in enumerate(dict_models):
            row.cells[i + 2].width = Inches(SIZE_METRIC_COL)
            prediction = dict_model['model'].predict_proba(X_val_var[[index]])[:, 1]
            predictions.append(prediction)
            row.cells[i + 2].paragraphs[0].add_run(str(np.round(prediction[0], 3)))
        row.cells[len(dict_models) + 2].width = Inches(SIZE_METRIC_COL)
        std = np.std(predictions)
        stds.append(std)
        row.cells[len(dict_models) + 2].paragraphs[0].add_run(str(np.round(std, 3)))

    report_doc.add_paragraph(f'The average standard deviation is {np.round(np.mean(stds),3)}')
    return


def write_paragraphs_dict(place, dict_to_write, font_size=SIZE_NORMAL_FONT, is_cell_table=False, bullet_level=1):
    # This first loop is for avoiding an empty first paragraph
    if is_cell_table:
        for i in range(len(dict_to_write.keys())-1):
            place.add_paragraph()

    for i, (param_name, param_value) in enumerate(dict_to_write.items()):
        if isinstance(param_value, float):
            valor = np.round(param_value, 3)
        else:
            valor = param_value

        if is_cell_table:
            paragraph = place.paragraphs[i]
        else:
            paragraph = place.add_paragraph()
        if bullet_level == 1:
            paragraph.style = 'List Bullet'
        else:
            paragraph.style = 'List Bullet ' + str(bullet_level)
        paragraph.add_run(f'{param_name}: {valor}').font.size = Pt(font_size)
    return


def print_search_spaces(report_doc, model_search_spaces):
    report_doc.add_paragraph(f'The search spaces for the optimization is the following:')
    for model, search_space in model_search_spaces.items():
        report_doc.add_paragraph(f'Search space for {model} model.').style = 'List Bullet'
        for key, value in search_space.items():
            if key != 'search_space':
                report_doc.add_paragraph(f'{key}: {value}').style = 'List Bullet 2'
            else:
                report_doc.add_paragraph(f'Search space: ').style = 'List Bullet 2'
                for item in value:
                    report_doc.add_paragraph(f'{item.name}: {item}').style = 'List Bullet 3'

def write_intro_doc(report_doc, y, model_search_spaces,
            k_outer_fold, skip_outer_folds, k_inner_fold,
            skip_inner_folds, n_initial_points, n_calls,
            calibrated, loss_metric, size_variance_validation,
            skopt_func):
    report_doc.add_heading('Introduction', level=1)
    report_doc.add_paragraph(f'Report of search and training made on {datetime.now().strftime("%B %d, %Y at %X")}.')
    report_doc.add_heading('Training data', level=2)
    p = report_doc.add_paragraph()
    p.add_run(f'There are {len(y)} training samples. ')
    p.add_run(f'The distribution of the labels is the following:')
    class_labels = Counter(y)
    for key in ordered_keys_classes(Counter(y)):
        p = report_doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(f'Class {key}: {class_labels[key]} instances.')
    report_doc.add_heading('Optimizing procedure', level=2)
    report_doc.add_paragraph('The parameters for the bayesian search are: ')
    report_doc.add_paragraph(f'Nested Cross Validation using {k_outer_fold} outer folds and '
                             f'{k_inner_fold} inner folds.').style = 'List Bullet'
    report_doc.add_paragraph(f'Some of the folds will be skipped. In particular, {skip_outer_folds} outer folds '
                             f'and {skip_inner_folds} inner folds will be skipped.').style = 'List Bullet'
    report_doc.add_paragraph(f'For each outer fold search, a model will be fitted. In order to search for the best '
                             f'hyperparameters, {n_initial_points} initial points will be evaluated, and '
                             f'{n_calls} additional calls will be made.').style = 'List Bullet'
    if calibrated:
        report_doc.add_paragraph(f'Models will be calibrated using their inner validation set.').style = 'List Bullet'
    else:
        report_doc.add_paragraph(f'Models will not be calibrated.').style = 'List Bullet'
    report_doc.add_paragraph(f'The optimizing metric for the bayesian search is {loss_metric}.').style = 'List Bullet'
    report_doc.add_paragraph(f'The function used for the bayesian search is {skopt_func.__name__}.').style = 'List Bullet'
    report_doc.add_paragraph(f'Additionally, {size_variance_validation} instances will be left out '
                             f'for assessing the variance of all models.').style = 'List Bullet'
    print_search_spaces(report_doc, model_search_spaces)
    return


def write_train_report(report_doc, list_params, list_metrics,
                       list_holdout_metrics, peeking_metrics, list_comments):
    print_holdout_metrics = len(list_holdout_metrics) > 0
    report_doc.add_heading(f'Report of training in this outer fold', level=2)
    list_loss_metrics = [metric['loss_metric'] for metric in list_metrics]
    index_best_model = list_loss_metrics.index(min(list_loss_metrics))
    report_doc.add_paragraph(
        f'Best model with respect to selected metric is {list_params[index_best_model]["model"]} with the following params:')
    write_paragraphs_dict(report_doc,
                          {key: value for (key, value) in list_params[index_best_model].items() if key != 'model'})

    report_doc.add_heading(f'Comparison of all models trained in this outer fold', level=3)
    # Create table for comparison
    if print_holdout_metrics:
        table = report_doc.add_table(rows=len(list_params) + 1, cols=len(peeking_metrics) * 2 + 3)
    else:
        table = report_doc.add_table(rows=len(list_params) + 1, cols=len(peeking_metrics) + 3)
    table.style = 'TableGrid'
    table.autofit = False
    # Fill first row of table
    first_row = table.rows[0]
    first_row.cells[0].width = Inches(SIZE_MODEL_COL)
    first_row.cells[0].paragraphs[0].add_run('Model').bold = True
    first_row.cells[1].width = Inches(SIZE_PARAM_COL)
    first_row.cells[1].paragraphs[0].add_run('Params').bold = True
    first_row.cells[2].width = Inches(SIZE_COMMENT_COL)
    first_row.cells[2].paragraphs[0].add_run('Comments').bold = True
    for i, metric in enumerate(peeking_metrics):
        first_row.cells[i + 3].width = Inches(SIZE_METRIC_COL)
        first_row.cells[i + 3].paragraphs[0].add_run(metric + ' on inner fold').bold = True
    if print_holdout_metrics:
        for i, metric in enumerate(peeking_metrics):
            first_row.cells[i + 3 + len(peeking_metrics)].width = Inches(SIZE_METRIC_COL)
            first_row.cells[i + 3 + len(peeking_metrics)].paragraphs[0].add_run(metric + ' on outer fold').bold = True
    # Fill other rows
    for row_index, params in enumerate(list_params):
        row = table.rows[row_index + 1]
        if index_best_model == row_index:
                color_row(row)
        row.cells[0].width = Inches(SIZE_MODEL_COL)
        row.cells[0].text = params['model']
        row.cells[1].width = Inches(SIZE_PARAM_COL)
        write_paragraphs_dict(row.cells[1],
                              {key: value
                              for (key, value) in list_params[row_index].items() if key != 'model'},
                              SIZE_SMALL_FONT,
                              is_cell_table=True)
        row.cells[2].width = Inches(SIZE_COMMENT_COL)
        write_paragraphs_dict(row.cells[2], list_comments[row_index], SIZE_SMALL_FONT, is_cell_table=True)
        for i, metric in enumerate(peeking_metrics):
            row.cells[i + 3].width = Inches(SIZE_METRIC_COL)
            row.cells[i + 3].text = str(np.round(list_metrics[row_index]['peeking_metrics'][metric], 3))
        if print_holdout_metrics:
            for i, metric in enumerate(peeking_metrics):
                row.cells[i + 3 + len(peeking_metrics)].width = Inches(SIZE_METRIC_COL)
                row.cells[i + 3 + len(peeking_metrics)].text = str(
                    np.round(list_holdout_metrics[row_index]['peeking_metrics'][metric], 3))
    return


def reporting_width(report_level, peeking_metrics):
    if report_level in [1, 11]:
        return max(8.5, len(peeking_metrics) * 2 * SIZE_METRIC_COL + SIZE_MODEL_COL + SIZE_PARAM_COL + SIZE_COMMENT_COL + 2)
    else:
        return max(8.5, len(peeking_metrics) * SIZE_METRIC_COL + SIZE_MODEL_COL + SIZE_PARAM_COL + SIZE_COMMENT_COL + 2)


def merge_docs(first_doc, second_doc):
    if not first_doc:
        return second_doc
    if not second_doc:
        return first_doc
    composer = Composer(first_doc)
    composer.append(second_doc)
    return composer.doc
