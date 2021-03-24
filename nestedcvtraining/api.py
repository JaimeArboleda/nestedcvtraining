import sklearn.pipeline
from sklearn.model_selection import StratifiedKFold, train_test_split
from skopt import gp_minimize
from docx import Document
from docx.shared import Inches
from nestedcvtraining.utils.training import train_inner_model
from nestedcvtraining.utils.metrics import is_supported
from nestedcvtraining.utils.reporting import evaluate_model, reporting_width, merge_docs, write_intro_doc
from collections import Counter
from sklearn.base import TransformerMixin
from skopt.space.transformers import Identity
from skopt.space import Dimension
from sklearn.utils.validation import check_X_y
from inspect import isfunction, signature




def find_best_binary_model(
        X,
        y,
        model_search_spaces,
        k_outer_fold=5,
        skip_outer_folds=[],
        k_inner_fold=5,
        skip_inner_folds=[],
        n_initial_points=5,
        n_calls=10,
        calibrated=True,
        loss_metric='average_precision',
        peeking_metrics=[],
        report_level=11,
        size_variance_validation=20,
        skopt_func=gp_minimize,
        verbose=False,
        build_final_model=True
    ):
    """Finds best binary calibrated classification model and optionally
    generate a report doing a nested cross validation. In the inner
    cross validation, doing a Bayesian Search, the best parameters are found.
    In the outer cross validation, the model is validated.
    Finally, the whole procedure is used for the full dataset to return
    the best possible model.


    Parameters
    ----------
    X : np.array
        Feature set.

    y : np.array
        Classification target to predict. For the moment only binary labels are allowed, and
        values are supposed to be {0, 1} or {-1, 1}

    model_search_spaces : Dict[str : List[List[skopt.Space]]
        Dict of models to try inside of the inner loops. For each model, there is
        the corresponding list of space objects to delimit where the parameters live,
        including the pipeline postprocess to make. It admits also an option to set
        undersampling_majority_class method. It admits two values, True or False. If True
        it builds an ensemble model in the inner loop by creating many balanced folds
        by using the minority class with a undersampling of the majority class. If using
        this option, it also admits an Int max_k_undersampling, in order to limit the number of
        splits made for this (because if the imbalance ratio is for example 1000:1, it will
        create 1000 splits, which can be too much).

    k_outer_fold : int, default=5
        Number of folds for the outer cross-validation.

    skip_outer_folds : list, default=None
        If set, list of folds to skip during the loop.

    k_inner_fold : int, default=5
        Number of folds for the inner cross-validation.

    skip_inner_folds : list, default=None
        If set, list of folds to skip during the loop.

    n_initial_points : int, default=5
        Number of initial points to use in Bayesian Optimization.

    n_calls : int, default=5
        Number of additional calls to use in Bayesian Optimization.

    calibrated : bool, default=False
        If True, all models are calibrated using CalibratedClassifierCV

    loss_metric : str, default='auc'
        Metric to use in order to find best parameters in Bayesian Search. Options:
        - roc_auc
        - average_precision
        - neg_brier_score
        - neg_log_loss
        - histogram_width

    peeking_metrics : List[str], default=[]
        If not empty, in the report there will be a comparison between the metric of
        evaluation on the inner fold and the list of metrics in peeking_metrics.

    report_levels : int, default=11
        If 00, no report is returned.
        If 01, plots are not included. All peeking-metrics are evaluated on the outer fold for each inner-fold model.
        If 10, plots are included. No evaluation of peeking-metrics on the outer fold for each inner-fold model.
        If 11, a full report (it can be more time consuming).

    size_variance_validation : int, default=20
        Number of samples to use to check variance of different models.

    skopt_func : callable, default=gp_minimize
        Minimization function of the skopt library to be used.

    verbose : bool, default=False
        If True, you can trace the progress in the terminal.

    build_final_model : bool, default=True
        If False, no final model is built (only the report doc is returned). It can be convenient
        during the experimental phase.

    Returns
    -------
    model : Model trained with the full dataset using the same procedure
    as in the inner cross validation.
    doc : Document python-docx if report_level > 0. Otherwise, None



        report_level=11,
    """

    # Validation of inputs
    X, y = check_X_y(X, y,
                     accept_sparse=['csc', 'csr', 'coo'],
                     force_all_finite=False, allow_nd=True)
    counter = Counter(y)
    if len(counter) > 2:
        raise NotImplementedError("Multilabel classification is not yet implemented")
    y_values = set(counter.keys())
    if y_values != {-1, 1} and y_values != {0, 1}:
        raise NotImplementedError("Values of target are expected to be in {0, 1} or in {-1, 1}")

    if not _validate_model_search_space(model_search_spaces):
        raise ValueError("model_search_spaces is not well formed")

    if not _validate_folds(k_outer_fold, skip_outer_folds, k_inner_fold, skip_inner_folds):
        raise ValueError("Folds parameters are not well formed")

    if not _validate_bayesian_search(n_initial_points, n_calls, skopt_func):
        raise ValueError("Bayesian search parameters are not well formed")

    if not is_supported(loss_metric):
        raise NotImplementedError(f"Loss metric {loss_metric} is not implemented.")

    if not isinstance(peeking_metrics, list):
        raise ValueError("Peeking metrics must be a list of str")

    for metric in peeking_metrics:
        if not is_supported(metric):
            raise NotImplementedError(f"Metric {metric} is not implemented.")

    if not isinstance(calibrated, bool):
        raise ValueError("calibrated must be a boolean")

    if not isinstance(verbose, bool):
        raise ValueError("verbose must be a boolean")

    if not isinstance(build_final_model, bool):
        raise ValueError("build_final_model must be a boolean")

    if not isinstance(size_variance_validation, int):
        raise ValueError("size_variance_validation must be an int")

    if size_variance_validation < 0 or size_variance_validation > len(y):
        raise ValueError("size_variance_validation cannot be negative nor bigger than number of instances")

    if report_level not in [0, 1, 10, 11]:
        raise ValueError("report_level must be either 0, 1, 10 or 11")

    # End of validation of inputs

    if loss_metric not in peeking_metrics:
        peeking_metrics.append(loss_metric)

    if report_level > 0:
        outer_report_doc = Document()
        section = outer_report_doc.sections[0]
        section.page_width = Inches(reporting_width(report_level, peeking_metrics))
        outer_report_doc.add_heading('Report of training', 0)

        write_intro_doc(
            outer_report_doc, y, model_search_spaces,
            k_outer_fold, skip_outer_folds, k_inner_fold,
            skip_inner_folds, n_initial_points, n_calls,
            calibrated, loss_metric, size_variance_validation,
            skopt_func)
        inner_report_doc = Document()
        section = inner_report_doc.sections[0]
        section.page_width = Inches(reporting_width(report_level, peeking_metrics))
    else:
        outer_report_doc = None
        inner_report_doc = None

    X, X_val_var, y, y_val_var = train_test_split(X, y, test_size=size_variance_validation, random_state=42, stratify=y)

    outer_cv = StratifiedKFold(n_splits=k_outer_fold)
    dict_inner_models = []
    outer_Xs = []
    outer_ys = []
    folds_index = []
    inner_report_doc.add_heading(f'Report of inner trainings', level=1)
    for k, (train_index, test_index) in enumerate(outer_cv.split(X, y)):
        if k not in skip_outer_folds:
            folds_index.append(k)
            inner_report_doc.add_heading(f'Report of inner training in fold {k} of outer Cross Validation', level=2)
            X_hold_out = X[test_index] if report_level in [1, 11] else []
            y_hold_out = y[test_index] if report_level in [1, 11] else []
            inner_model, model_params, model_comments = train_inner_model(
                X=X[train_index], y=y[train_index], model_search_spaces=model_search_spaces,
                X_hold_out=X_hold_out, y_hold_out=y_hold_out,
                k_inner_fold=k_inner_fold, skip_inner_folds=skip_inner_folds,
                n_initial_points=n_initial_points, n_calls=n_calls,
                calibrated=calibrated, loss_metric=loss_metric, peeking_metrics=peeking_metrics,
                verbose=verbose, skopt_func=skopt_func, report_doc=inner_report_doc)
            dict_inner_models.append({'model': inner_model,
                                 'params': model_params,
                                 'comments': model_comments})

            outer_Xs.append(X[test_index])
            outer_ys.append(y[test_index])
    outer_report_doc.add_heading(f'Report of validation of the model in the outer Cross Validation', level=1)
    add_plots = True if report_level > 9 else False
    evaluate_model(
        dict_models=dict_inner_models, Xs=outer_Xs, ys=outer_ys,
        X_val_var=X_val_var, y_val_var=y_val_var,
        folds_index=folds_index, report_doc=outer_report_doc,
        loss_metric=loss_metric, peeking_metrics=peeking_metrics,
        add_plots=add_plots
    )
    # After assessing the procedure, we repeat it on the full dataset:
    final_model = None
    if build_final_model:
        final_model, _, _ = train_inner_model(
                X=X, y=y, model_search_spaces=model_search_spaces,
                X_hold_out=[], y_hold_out=[],
                k_inner_fold=k_inner_fold, skip_inner_folds=skip_inner_folds,
                n_initial_points=n_initial_points, n_calls=n_calls,
                calibrated=calibrated, loss_metric=loss_metric, peeking_metrics=[],
                verbose=verbose, skopt_func=skopt_func, report_doc=None)
    return final_model, merge_docs(outer_report_doc, inner_report_doc)


class OptionedPostProcessTransformer(TransformerMixin):
    """This class converts a dictionary of different options of post-process
    (each option will be a fully defined pipeline) in a transformer that
    has only one parameter, the option.
    In this way, you can perform a bayesian optimization search including a postprocess
    pipeline that uses this set of categorical fixed options.

    Example of input dict:
        dict_pipelines_post_process = {
        "option_1": Pipeline(
            [("scale", StandardScaler()), ("reduce_dims", PCA(n_components=5))]
        ),
        "option_2": Pipeline(
            [
                ("scale", StandardScaler()),
                ("reduce_dims", SelectKBest(mutual_info_classif, k=5)),
            ]
        ),
        "option_3": Pipeline([("identity", MidasIdentity())])
        }
    """

    def __init__(self, dict_pipelines):
        if not self._validate_dict(dict_pipelines):
            raise ValueError("The dictionary is not well formed. Please check the docs and examples. ")
        if not self._validate_resamplers(dict_pipelines):
            raise ValueError("OptionedPostProcessTransformer does not "
                             "support resamplers inside ")
        self.dict_pipelines = dict_pipelines
        self.option = list(dict_pipelines.keys())[0]
        super().__init__()

    def fit(self, X, y=None):
        self.dict_pipelines[self.option].fit(X, y)
        return self

    def set_params(self, **params):
        self.option = params['option']
        return self

    def transform(self, X):
        return self.dict_pipelines[self.option].transform(X)

    def fit_transform(self, X, y=None):
        return self.dict_pipelines[self.option].fit_transform(X, y)

    def _validate_dict(self, dict_pipelines):
        check_pipelines = any(
            [
                not isinstance(pipeline, sklearn.pipeline.Pipeline)
                for pipeline in dict_pipelines.values()
            ]
        )
        return not check_pipelines

    def _validate_resamplers(self, dict_pipelines):
        all_steps = []
        for pipeline in dict_pipelines.values():
            all_steps.extend(pipeline.steps)
        exists_resampler = any(
            [
                hasattr(step[1], "fit_resample")
                for step in all_steps
            ]
        )
        return not exists_resampler


class MidasIdentity(Identity):
    """
    Convenience class for creating a Pipeline that does not perform any transformation.
    It can be handy when in combination with OptionedPostProcessTransformer.
    """
    def fit(self, X, y):
        return self


def _validate_model_search_space(model_search_spaces):
    if not isinstance(model_search_spaces, dict):
        raise ValueError("model_search_spaces must be a dict")
    for model_search in model_search_spaces.values():
        if not isinstance(model_search, dict):
            raise ValueError("model_search_spaces must be a dict of dicts")
        for (key, value) in model_search.items():
            if key not in ['model', 'pipeline_post_process', 'search_space']:
                raise ValueError("Some inner key is not in ['model', 'pipeline_post_process', 'search_space']")
            if key == 'model':
                if not hasattr(value, 'predict_proba'):
                    raise ValueError("Estimator has not predict_proba method")
                if not hasattr(value, 'fit'):
                    raise ValueError("Estimator has not fit method")
            if key == 'pipeline_post_process':
                if value:
                    if not isinstance(value, sklearn.pipeline.Pipeline):
                        raise ValueError("pipeline_post_process is not a pipeline")
            if key == 'search_space':
                if not isinstance(value,list):
                    raise ValueError("search_space is not a list")
                for element in value:
                    if not isinstance(element, Dimension):
                        raise ValueError("search_space is not valid")
    return True


def _validate_folds(k_outer_fold, skip_outer_folds, k_inner_fold, skip_inner_folds):
    if not isinstance(k_outer_fold, int):
        raise ValueError("k_outer_fold must be int")
    if not isinstance(k_inner_fold, int):
        raise ValueError("k_inner_fold must be int")
    if not isinstance(skip_outer_folds, list):
        raise ValueError("skip_outer_folds must be a list")
    if not isinstance(skip_inner_folds, list):
        raise ValueError("k_outer_fold must be a list")
    for element in skip_outer_folds:
        if element not in list(range(k_outer_fold)):
            raise ValueError("skip_outer_folds must be contained in [0, k_outer_folds-1]")
    for element in skip_inner_folds:
        if element not in list(range(k_inner_fold)):
            raise ValueError("skip_inner_folds must be contained in [0, k_inner_fold-1]")
    return True


def _validate_bayesian_search(n_initial_points, n_calls, skopt_func):
    if not isinstance(n_initial_points, int):
        raise ValueError("n_initial_points must be int")
    if not isinstance(n_calls, int):
        raise ValueError("n_calls must be int")
    if n_initial_points < 0 or n_calls < 0:
        raise ValueError("n_calls and n_initial_points must positive")
    if n_calls < n_initial_points:
        raise ValueError("n_calls cannot be lesser than n_initial_points")
    if not isfunction(skopt_func):
        raise ValueError("skopt_func must be callable")
    signature_skopt_func = signature(skopt_func)
    kwargs = [parameter for parameter in signature_skopt_func.parameters]
    if 'func' not in kwargs:
        raise ValueError("skopt_func must have func in kwargs")
    if 'dimensions' not in kwargs:
        raise ValueError("skopt_func must have dimensions in kwargs")
    if 'n_initial_points' not in kwargs:
        raise ValueError("skopt_func must have n_initial_points in kwargs")
    if 'n_calls' not in kwargs:
        raise ValueError("skopt_func must have n_calls in kwargs")
    return True

