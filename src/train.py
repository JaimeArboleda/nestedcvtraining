from sklearn.model_selection import StratifiedKFold, train_test_split
from skopt import gp_minimize
from docx import Document
from docx.shared import Inches
from utils.training import train_inner_calibrated_binary_model
from utils.reporting import evaluate_model, reporting_width, merge_docs, write_intro_doc











def find_best_binary_model(
        X,
        y,
        model_search_spaces,
        k_outer_fold=5,
        skip_outer_folds=[],
        k_inner_fold=5,
        skip_inner_folds=[],
        n_initial_points=5,
        n_calls=10,
        calibrated=True,
        loss_metric='average_precision',
        peeking_metrics=[],
        report_level=11,
        size_variance_validation=20,
        skopt_func=gp_minimize,
        verbose=False,
        build_final_model=True
    ):
    """Finds best binary calibrated classification model and optionally
    generate a report doing a nested cross validation. In the inner
    cross validation, doing a Bayesian Search, the best parameters are found.
    In the outer cross validation, the model is validated.
    Finally, the whole procedure is used for the full dataset to return
    the best possible model.


    Parameters
    ----------
    X : np.array
        Feature set.

    y : np.array
        Classification target to predict.

    model_search_spaces : Dict[str : List[List[skopt.Space]]
        Dict of models to try inside of the inner loops. For each model, there is
        the corresponding list of space objects to delimit where the parameters live,
        including the pipeline postprocess to make. It admits also an option to set
        undersampling_majority_class method. It admits two values, True or False. If True
        it builds an ensemble model in the inner loop by creating many balanced folds
        by using the minority class with a undersampling of the majority class. If using
        this option, it also admits an Int max_k_undersampling, in order to limit the number of
        splits made for this (because if the imbalance ratio is for example 1000:1, it will
        create 1000 splits, which can be too much).

    k_outer_fold : int, default=5
        Number of folds for the outer cross-validation.

    skip_outer_folds : list, default=None
        If set, list of folds to skip during the loop.

    k_inner_fold : int, default=5
        Number of folds for the inner cross-validation.

    skip_inner_folds : list, default=None
        If set, list of folds to skip during the loop.

    n_initial_points : int, default=5
        Number of initial points to use in Bayesian Optimization.

    n_calls : int, default=5
        Number of additional calls to use in Bayesian Optimization.

    calibrated : bool, default=False
        If True, all models are calibrated using CalibratedClassifierCV

    loss_metric : str, default='auc'
        Metric to use in order to find best parameters in Bayesian Search. Options:
        - roc_auc
        - average_precision
        - neg_brier_score
        - neg_log_loss
        - histogram_width

    peeking_metrics : List[str], default=[]
        If not empty, in the report there will be a comparison between the metric of
        evaluation on the inner fold and the list of metrics in peeking_metrics.

    report_levels : int, default=11
        If 00, no report is returned.
        If 01, plots are not included. All peeking-metrics are evaluated on the outer fold for each inner-fold model.
        If 10, plots are included. No evaluation of peeking-metrics on the outer fold for each inner-fold model.
        If 11, a full report (it can be more time consuming).

    size_variance_validation : int, default=20
        Number of samples to use to check variance of different models.

    skopt_func : callable, default=gp_minimize
        Minimization function of the skopt library to be used.

    verbose : bool, default=False
        If True, you can trace the progress in the terminal.

    build_final_model : bool, default=True
        If False, no final model is built (only the report doc is returned). It can be convenient
        during the experimental phase.

    Returns
    -------
    model : Model trained with the full dataset using the same procedure
    as in the inner cross validation.
    doc : Document python-docx if report_level > 0. Otherwise, None
    """
    # TODO: Check inputs
    if loss_metric not in peeking_metrics:
        peeking_metrics.append(loss_metric)

    if report_level > 0:
        outer_report_doc = Document()
        section = outer_report_doc.sections[0]
        section.page_width = Inches(reporting_width(report_level, peeking_metrics))
        outer_report_doc.add_heading('Report of training', 0)

        write_intro_doc(
            outer_report_doc, y, model_search_spaces,
            k_outer_fold, skip_outer_folds, k_inner_fold,
            skip_inner_folds, n_initial_points, n_calls,
            calibrated, loss_metric, size_variance_validation,
            skopt_func)
        inner_report_doc = Document()
        section = inner_report_doc.sections[0]
        section.page_width = Inches(reporting_width(report_level, peeking_metrics))
    else:
        outer_report_doc = None
        inner_report_doc = None

    X, X_val_var, y, y_val_var = train_test_split(X, y, test_size=size_variance_validation, random_state=42, stratify=y)

    outer_cv = StratifiedKFold(n_splits=k_outer_fold)
    dict_inner_models = []
    outer_Xs = []
    outer_ys = []
    folds_index = []
    inner_report_doc.add_heading(f'Report of inner trainings', level=1)
    for k, (train_index, test_index) in enumerate(outer_cv.split(X, y)):
        if k not in skip_outer_folds:
            folds_index.append(k)
            inner_report_doc.add_heading(f'Report of inner training in fold {k} of outer Cross Validation', level=2)
            inner_model, model_params, model_comments = train_inner_calibrated_binary_model(
                X=X[train_index], y=y[train_index], k_inner_fold=k_inner_fold, skip_inner_folds=skip_inner_folds,
                X_hold_out=X[test_index], y_hold_out=y[test_index],
                report_doc=inner_report_doc, n_initial_points=n_initial_points,
                n_calls=n_calls,
                dict_model_params=model_search_spaces,
                loss_metric=loss_metric, peeking_metrics=peeking_metrics, verbose=verbose,
                skopt_func=skopt_func)
            dict_inner_models.append({'model': inner_model,
                                 'params': model_params,
                                 'comments': model_comments})

            outer_Xs.append(X[test_index])
            outer_ys.append(y[test_index])
    outer_report_doc.add_heading(f'Report of validation of the model in the outer Cross Validation', level=1)
    evaluate_model(
        dict_models=dict_inner_models, Xs=outer_Xs, ys=outer_ys, X_val_var=X_val_var, y_val_var=y_val_var,
        folds_index=folds_index, report_doc=outer_report_doc, loss_metric=loss_metric, peeking_metrics=peeking_metrics
    )
    # After assessing the procedure, we repeat it on the full dataset:
    final_model = None
    if build_final_model:
        final_model, _, _ = train_inner_calibrated_binary_model(
                X=X, y=y, k_inner_fold=k_inner_fold, skip_inner_folds=skip_inner_folds,
                report_doc=None, n_initial_points=n_initial_points,
                n_calls=n_calls,
                dict_model_params=model_search_spaces,
                loss_metric=loss_metric, verbose=verbose, skopt_func=skopt_func)
    return final_model, merge_docs(outer_report_doc, inner_report_doc)














